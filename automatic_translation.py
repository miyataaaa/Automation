from docx import Document
import os 
import chromedriver_binary 
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
import time
from threading import Lock
import time
from concurrent.futures import ThreadPoolExecutor
import math 
options = Options()
options.add_argument('--disable-gpu')
options.add_argument('--disable-extensions')
options.add_argument('--proxy-server="direct://"')
options.add_argument('--proxy-bypass-list=*')
options.add_argument('--start-maximized')
lock = Lock()


def sentense_from_docx(**kwargs):
    
    """
    翻訳対象のPDFファイルから、翻訳したい部分を手動でコピペしたwordファイルから英文を抜き出して
    1つのリストにする関数。
    ファイル内の構造は1つの2列×n行テーブルを持ち、左側の列のセルに翻訳したい塊ごとに英文が入っている事が前提
    """
    
    parent_path = kwargs['Fpath']
    fname = kwargs['Fname'] + ".docx"
    fpath = os.path.join(parent_path, fname)
    docxfile = Document(fpath)
    tables = docxfile.tables[0]
    column = tables.columns[0]
    sentenses = []
    for cell in column.cells:
        text = cell.text
#         print(text)
        sentenses.append(text)
        
    return sentenses

def translation_deepl(sourse_texts=[]):
    
    """
    抜き出した英文をdeeplにコピペして翻訳する関数。引数は翻訳する文章が要素のリスト
    """
    browser = webdriver.Chrome()
    url = 'https://www.deepl.com/ja/translator'
    browser.get(url)
    # deepleにアクセスするまでしばらく待つ
    time.sleep(1)
    # wordの文章をparagraph単位で翻訳していく
    # len(sourse_texts)
    translated_texts = []
    for i in range(len(sourse_texts)):         
        sourse_text = sourse_texts[i]
        if len(sourse_text) < 10:
            print("i={} continue\n: {}".format(i, sourse_text))
            translated_texts.append("")
#             print("continue\n")
            continue
        else:
            print("i={}: {}".format(i, sourse_text))
            
        stextarea = browser.find_element_by_css_selector(
            '.lmt__textarea.lmt__source_textarea.lmt__textarea_base_style')
        ttextarea = browser.find_element_by_css_selector(
            '.lmt__textarea.lmt__target_textarea.lmt__textarea_base_style')

        lock.acquire()
        stextarea.send_keys(sourse_text)
        lock.release()
        #time.sleep(2)

        translated_text = ''

        lock.acquire()
        while not translated_text:
            time.sleep(6)
            translated_text = ttextarea.get_property('value')
            translated_texts.append(translated_text)
        # ここに翻訳結果が出力されたら成功（途中経過をjupyter notebookで確認する場合はコメントイン）
        #print(translated_text)    
        lock.release()

        # sourse_textareaをクリアする
        time.sleep(2)
        stextarea.send_keys(Keys.CONTROL, "a")
        stextarea.send_keys(Keys.BACKSPACE)
        
    return translated_texts
    
def split_sentenses(max_worker, sentenses=[]):
    
    """
    並列処理の前処理として翻訳する英文のリストを並列処理に使用する論理プロセッサの数で分割する関数
    """
    
    n = math.ceil(len(sentenses)/max_worker)
    split_sentenses = [sentenses[i:i+n] for i in range(0, len(sentenses), n)]
    
    return split_sentenses

def run_multiThread(sentenses=[], **kwargs):
    
    """
    translation_deepl関数を、引数辞書で指定した論理プロセッサ数で処理する。
    """
    max_worker = kwargs['max_worker']
    split_list = split_sentenses(max_worker, sentenses=sentenses)
    print("len(split_list): {}".format(len(split_list)))
    for i in range(len(split_list)):
        texts = split_list[i]
        print("i={}\nsplit_list: {}".format(i, texts))
    
    with ThreadPoolExecutor(max_workers=max_worker) as executor:
        
        result = executor.map(translation_deepl, split_list)
    
    translated_texts = []
    results = [sentenses for sentenses in result]
    for i in range(len(results)):
        texts = results[i]
#         print("len(translated_texts{}): {}\n".format(i, len(texts)))
        for j in range(len(texts)):
#             print("i={} j={}\n{}".format(i, j, texts[j]))
            translated_texts.append(texts[j])
    
    return translated_texts


def sentense_to_docx(sentenses, **kwargs):
    
    parent_path = kwargs['Fpath']
    fname = kwargs['Fname'] + ".docx"
    fpath = os.path.join(parent_path, fname)
    docxfile = Document(fpath)
    tables = docxfile.tables[0]
    cells = tables.column_cells(1)
    
    for i in range(len(sentenses)):
        cells[i].text = sentenses[i]
        
    docxfile.save(fpath)

def main(**kwargs):
    
    sentenses = sentense_from_docx(**kwargs)
    translated_texts = run_multiThread(sentenses=sentenses, **kwargs)
    sentense_to_docx(translated_texts, **kwargs)
    
if __name__ == "__main__":
    
    kwargs = {"Fpath": r"C:\Users\miyar\OneDrive\デスクトップ\四年\文献演習",
          "Fname": r"test_Timescales of landscape response to divide migration",
          "max_worker": 4}
    
    main(**kwargs)